from flask import (
    Flask,
    render_template,
    request,
    session,
    g,
    redirect,
    url_for,
    jsonify,
    json
)
from flask_cors import CORS
from pymongo import MongoClient
from flask_pymongo import PyMongo
import datetime
import hashlib
import pymongo
import traceback
import sys
import bcrypt
import json


import random
from time import time
from uuid import uuid4
# import requests
# pip3 install python-docx
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from random import SystemRandom

# 勇----------------------------------------------------------------------------------------------------
cluster = MongoClient(
    "mongodb+srv://admin:YpjVk0DElOtbcLda@cluster.8ae9h.mongodb.net/EcoChain?retryWrites=true&w=majority")
db = cluster["EcoChain"]
colChain = db["Chain"]
colRecChain = db["RecChain"]
colQueue = db["Queue"]
colCustomer = db["customer"]
db2 = cluster["Member"]
alphabet = u'9ABCDEFGHIJKLMNOPQRSTUVWXYZ'
generator = SystemRandom()


def Compare(li1, li2):
    return (list(set(li1).symmetric_difference(set(li2))))


class Blockchain:
    def __init__(self):
        self.current_transactions = []
        #self.chain = []
        self.chain = list(colChain.find({}))
        self.record_transactions = ''
        self.key = 119
        RecChain = colRecChain.find_one({})
        if RecChain == None:
            self.recChain = []
        else:
            self.recChain = list(RecChain['content'])
        self.perfectList = []
        # Create the genesis block
        if self.chain == []:
            self.new_block(previous_hash='1', proof=100)

    # 新區塊
    def new_block(self, proof, previous_hash):
        """
        Create a new Block in the Blockchain
        :param proof: The proof given by the Proof of Work algorithm
        :param previous_hash: Hash of previous Block
        :return: New Block
        """

        block = {
            '_id': len(self.chain) + 1,
            'timestamp': time(),
            'transactions': self.current_transactions,
            'proof': proof,
            'previous_hash': previous_hash or self.hash(self.chain[-1]),
        }

        block2 = self.record_transactions
        # Reset the current list of transactions
        self.current_transactions = []
        self.record_transactions = ''

        self.chain.append(block)
        self.recChain.append(block2)

        return block

    # 新交易
    def new_transaction(self, sender, recipient, amount):
        """
        Creates a new transaction to go into the next mined Block
        :param sender: Address of the Sender
        :param recipient: Address of the Recipient
        :param amount: Amount
        :return: The index of the Block that will hold this transaction
        """
        print('len:  '+str(len(self.chain) + 1))
        print('sender:  '+str(sender))
        print('recipient:  '+str(recipient))
        print('amount:  '+str(amount))
        print(self.chain)
        # file=open("C:/Users/User/Desktop/PythonBlockchain/nodedata/OutPut.txt",'a')
        #temp='\n'+str(len(self.chain) + 1)+' '+str(sender)+' '+str(recipient)+' '+str(amount)
        record = str(len(self.chain) + 1)+' '+str(sender) + \
            ' '+str(recipient)+' '+str(amount)
        # file.write(temp)
        # file.close()

        self.record_transactions = record

        self.current_transactions.append({
            'sender': sender,
            'recipient': recipient,
            'amount': amount,
        })

        return self.last_block['_id'] + 1

    def insertData(self):
        colChain.remove({})
        colRecChain.remove({})
        colChain.insert_many(self.chain)
        colRecChain.insert_one({'_id': 1, 'content': self.recChain})

    def change(self):
        '''
        self.chain
        self.recChain
        block = {
            'index': len(self.chain) + 1,
            'timestamp': time(),
            'transactions': self.current_transactions,
            'proof': proof,
            'previous_hash': previous_hash or self.hash(self.chain[-1]),
        }
        '''
        self.certification()
        # temp=[2,'寄件者','收件者','金額']
        unionList = Compare(self.perfectList, self.recChain)
        print('與鏈上有不同之處:'+str(unionList))
        if unionList != []:
            for i in range(0, len(unionList)):
                if unionList[i] in self.perfectList and unionList[i] != '\n':
                    temp = unionList[i]
                    tempStr = temp.split(' ')
                    data1 = self.chain[int(tempStr[0])-1]
                    transactions = {
                        'sender': tempStr[1],
                        'recipient': tempStr[2],
                        'amount': tempStr[3].replace('\n', ''),
                    }
                    data1['transactions'] = transactions
                    if int(tempStr[0])+1 <= len(self.chain):
                        data2 = self.chain[int(tempStr[0])]
                        data2['previous_hash'] = self.hash(
                            self.chain[int(tempStr[0])])
                    record = str(tempStr[0])+' '+str(tempStr[1]) + \
                        ' '+str(tempStr[2])+' '+str(tempStr[3])
                    self.recChain[int(tempStr[0])-1] = record

    def certification(self):
        allCollections = db2.list_collection_names()
        ProfileCount = len(allCollections)
        print('共有'+str(ProfileCount)+'個檔案')
        ProfileNumArr = []
        WrongGetOut = 0
        count = 0
        temp = 1

        for i in range(0, ProfileCount+1):
            ProfileNumArr.append(i)
            #print('txtNumArr:'+str(txtNumArr[i])+' -- i:'+str(i))

        self.findPerfectList(ProfileCount, ProfileNumArr,
                             WrongGetOut, count, temp)
        print('結果--------------------------------')
        print('perfectList:'+str(self.perfectList))
        print()

    def findPerfectList(self, ProfileCount, ProfileNumArr, WrongGetOut, count, temp):
        while True:
            ProfileNum = random.randint(1, ProfileCount)
            allCollections = db2.list_collection_names()

            if ProfileNum in ProfileNumArr:
                print('第'+str(temp)+'輪--------------------------------')
                print('使用Member'+str(ProfileNum))
                WrongGetOut = 0
                count = 0
                MemberName = 'Mr'+str(ProfileNum)
                colMember = db2[self.hash(MemberName)]
                colMemberData = colMember.find_one({})
                dataMain = list(colMemberData['content'])

                for info in allCollections:
                    if info != MemberName:
                        ProfileMinor = db2[info]
                        ProfileMinorData = ProfileMinor.find_one({})
                        dataMinor = list(ProfileMinorData['content'])

                        if count == float(ProfileCount/2):
                            print('最裏頭 Count:'+str(count)+' -- 半數為 ' +
                                  str(float(ProfileCount)/2))
                            self.perfectList = dataMain
                            return
                        elif Compare(dataMain, dataMinor) == []:
                            count += 1
                            print('最裏頭 Count:'+str(count)+' -- 半數為 ' +
                                  str(float(ProfileCount)/2))
                        elif WrongGetOut == float(ProfileCount/2):
                            print('錯誤到半數跳過')
                            break
                        elif Compare(dataMain, dataMinor) != []:
                            WrongGetOut += 1
                            print('比對不同次數為 '+str(WrongGetOut))

                # 最後一筆才為半數所使用
                if count == float(ProfileCount/2):
                    print('最裏頭 Count:'+str(count)+' -- 半數為 ' +
                          str(float(ProfileCount)/2))
                    self.perfectList = dataMain
                    return

                ProfileNumArr.remove(ProfileNum)
                temp += 1
                print()

    def sendEmail(self, sender, recipient, amount, Ca):
        senderData = colCustomer.find_one({'_id': int(sender)})
        recipientData = colCustomer.find_one({'wallet': str(recipient)})
        senderEmail = str(senderData['email'])
        recipientEmail = str(recipientData['email'])
        senderUserName = str(senderData['username'])
        recipientUserName = str(recipientData['username'])

        gmail_user = 'EcoChainProjectOfficial@gmail.com'
        gmail_password = 'bestriventw'
        from_address = gmail_user

        # 寄給sender
        to_address = [senderEmail]
        Subject = "Successful Transaction"
        contents = "您與"+recipientUserName+"的交易成功。\n這次交易將扣除您"+str(amount)+"的金額和增加"+str(
            Ca)+"的度數，如有問題請務必寄信給客服尋求幫助。\n再次感謝您的購買與使用。\n\n\n搜尋碼:OFNGZHECEWTXPLTPJFGDEKHHULCVLJNSYBFD9FCKUJJA99LK9NDSTPY9CXJFPWUJQZIHTWCGZZQNJFXHB"
        mail = MIMEMultipart()
        mail['From'] = from_address
        mail['To'] = ', '.join(to_address)
        mail['Subject'] = Subject
        mail.attach(MIMEText(contents))
        smtpserver = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        smtpserver.ehlo()
        smtpserver.login(gmail_user, gmail_password)
        smtpserver.sendmail(from_address, to_address, mail.as_string())
        smtpserver.quit()

        # 寄給recipient
        to_address = [recipientEmail]
        Subject = "Successful Transaction"
        contents = "您與"+senderUserName+"的交易成功。\n這次交易您將獲得"+str(amount)+"的金額和扣除"+str(
            Ca)+"的度數，如有問題請務必寄信給客服尋求幫助。\n再次感謝您的購買與使用。\n\n\n搜尋碼:OFNGZHECEWTXPLTPJFGDEKHHULCVLJNSYBFD9FCKUJJA99LK9NDSTPY9CXJFPWUJQZIHTWCGZZQNJFXHB"
        mail = MIMEMultipart()
        mail['From'] = from_address
        mail['To'] = ', '.join(to_address)
        mail['Subject'] = Subject
        mail.attach(MIMEText(contents))
        smtpserver = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        smtpserver.ehlo()
        smtpserver.login(gmail_user, gmail_password)
        smtpserver.sendmail(from_address, to_address, mail.as_string())
        smtpserver.quit()

    def clientLogIn(self, userID):
        MemberName = 'Mr'+str(userID)
        colMember = db2[self.hash(MemberName)]
        colMemberData = colMember.find_one({})
        MemberRecChain = list(colMemberData['content'])
        root = 'C:/Users/'+self.hash(MemberName)+'.docx'

        try:
            self.Decrypt(root, self.key)
            doc = Document(root)
            dataWord = ''.join(str(e.text) for e in doc.paragraphs)
            datas = dataWord.split('\n')
            diff = Compare(MemberRecChain, datas)
            needModi = []

            for i in range(0, len(diff)):
                if diff[i] in datas:
                    needModi.append(diff[i])

            if needModi != []:
                for i in range(0, len(needModi)):
                    address = needModi[i].split(' ')

                    if int(address[0]) > len(MemberRecChain):
                        print('此資料庫已遭竄改')
                        return
                    else:
                        MemberRecChain[int(address[0])-1] = needModi[i]

        except Exception as e:
            print(e)

        self.Encrypt(root, self.key)
        colMember.remove({})
        colMember.insert_one({'_id': 1, 'content': MemberRecChain})

    def clientLogOut(self, userID):
        self.change()
        MemberName = 'Mr'+str(userID)
        dataStr = '\n'.join(str(e) for e in self.recChain)
        document = Document()
        document.add_paragraph(dataStr)
        # root='C:/Users/User/Desktop/PythonBlockchain/'+self.hash(MemberName)+'.docx'
        root = 'C:/Users/'+self.hash(MemberName)+'.docx'
        document.save(root)
        self.Encrypt(root, self.key)

    def clientRegister(self, userID):
        MemberName = 'Mr'+str(userID)
        colMember = db2[self.hash(MemberName)]
        colMember.remove({})
        colMember.insert_one({'_id': 1, 'content': self.recChain})

    def Encrypt(self, filename, key):
        file = open(filename, 'rb')
        data = file.read()
        file.close()

        data = bytearray(data)
        for index, value in enumerate(data):
            data[index] = value ^ key

        file = open(filename, 'wb')
        file.write(data)
        file.close()

    def Decrypt(self, filename, key):
        file = open(filename, 'rb')
        data = file.read()
        file.close()

        data = bytearray(data)
        for index, value in enumerate(data):
            data[index] = value ^ key

        file = open(filename, 'wb')
        file.write(data)
        file.close()
    '''
    def AllChange(self):
        self.change()
        allCollections=db2.list_collection_names()
        for i in allCollections:
            print(i)
            colMember=db2[i]
            #colMember.remove({})
            colMember.insert_one({'_id':1,'content':self.recChain})
     '''
# ----------------------------------------------------------------------先不用看
    @property
    def last_block(self):
        return self.chain[-1]

    @staticmethod
    def hash(block):
        """
        Creates a SHA-256 hash of a Block
        :param block: Block
        """

        # We must make sure that the Dictionary is Ordered, or we'll have inconsistent hashes
        block_string = json.dumps(block, sort_keys=True).encode()
        return hashlib.sha256(block_string).hexdigest()

    def proof_of_work(self, last_block):
        """
        Simple Proof of Work Algorithm:
         - Find a number p' such that hash(pp') contains leading 4 zeroes
         - Where p is the previous proof, and p' is the new proof

        :param last_block: <dict> last Block
        :return: <int>
        """

        last_proof = last_block['proof']
        last_hash = self.hash(last_block)

        proof = 0
        while self.valid_proof(last_proof, proof, last_hash) is False:
            proof += 1

        return proof

    @staticmethod
    def valid_proof(last_proof, proof, last_hash):
        """
        Validates the Proof
        :param last_proof: <int> Previous Proof
        :param proof: <int> Current Proof
        :param last_hash: <str> The hash of the Previous Block
        :return: <bool> True if correct, False if not.
        """

        guess = f'{last_proof}{proof}{last_hash}'.encode()
        guess_hash = hashlib.sha256(guess).hexdigest()
        return guess_hash[:1] == "0"
# 勇----------------------------------------------------------------------------------------------------


class User:
    def __init__(self, id, username, password):
        self.id = id
        self.username = username
        self.password = password

    def __repr__(self):
        return f'<User : {self.username}>'


app = Flask(__name__, static_url_path='', static_folder='../React')
CORS(app)
cluster = MongoClient(
    "mongodb+srv://admin:YpjVk0DElOtbcLda@cluster.8ae9h.mongodb.net/EcoChain?retryWrites=true&w=majority")
db = cluster["EcoChain"]
collection = db["customer"]

collection2 = db["Project2"]
collection3 = db["Image"]

app.secret_key = 'thissecretkeyonlyiknow'
file = '3碼郵遞區號.json'
with open(file, 'r') as obj:
    data = json.load(obj)
keyList = ["中心點經度", "中心點緯度", "TGOS_URL"]
for k in data:
    k["郵遞區號"] = k["_x0033_碼郵遞區號"]
    del k["_x0033_碼郵遞區號"]
for k in data:
    del k["中心點經度"]
for k in data:
    del k["中心點緯度"]
for k in data:
    del k["TGOS_URL"]

# 勇----------------------------------------------------------------------------------------------------
node_identifier = str(uuid4()).replace('-', '')

# Instantiate the Blockchain
blockchain = Blockchain()

# 開挖 ---------------------------------------------------
@app.route('/mine', methods=['GET'])
def mine():
    # We run the proof of work algorithm to get the next proof...
    last_block = blockchain.last_block
    proof = blockchain.proof_of_work(last_block)

    previous_hash = blockchain.hash(last_block)
    block = blockchain.new_block(proof, previous_hash)

    response = {
        'message': "New Block Forged",
        '_id': block['_id'],
        'transactions': block['transactions'],
        'proof': block['proof'],
        'previous_hash': block['previous_hash'],
    }
    return jsonify(response), 200

# 放交易 ---------------------------------------------------
@app.route('/transactions/new', methods=['POST'])
def new_transaction():
    values = json.loads(request.data)

    # Create a new Transaction
    _id = blockchain.new_transaction(
        values['sender'], values['recipient'], values['amount'])
    mine()
    blockchain.insertData()
    response = {'message': f'Transaction will be added to Block {_id}'}
    return jsonify(response), 201

# 整條鏈 ---------------------------------------------------
@app.route('/chain', methods=['GET'])
def full_chain():
    response = {
        'chain': blockchain.chain,
        'length': len(blockchain.chain),
    }
    print(blockchain.recChain)

    return jsonify(response), 200

# 更改資料---------------------------------------------------
@app.route('/change', methods=['GET'])
def change():
    blockchain.change()

    response = {
        'chain': blockchain.chain,
        'length': len(blockchain.chain),
    }

    return jsonify(response), 200


@app.route('/insertData', methods=['GET'])
def insertData():
    blockchain.insertData()
    response = {
        'chain': blockchain.chain,
        'length': len(blockchain.chain),
    }
    return jsonify(response), 200


@app.route('/email', methods=['POST'])
def sendEmail():
    values = json.loads(request.data)

    blockchain.sendEmail(
        values['sender'], values['recipient'], values['amount'], values['Ca'])
    response = {'message': f'Send Email !!'}
    return jsonify(response), 201


@app.route('/clientlogin', methods=['POST'])
def clientLogIn():

    #values = request.get_json()
    values = json.loads(request.data)

    blockchain.clientLogIn(values['userID'])
    response = {'message': f'Client Certification'}
    return jsonify(response), 201


@app.route('/clientlogout', methods=['POST'])
def clientLogOut():
    values = json.loads(request.data)
    blockchain.clientLogOut(values['userID'])
    response = {'message': f'Client Log Out'}
    return jsonify(response), 201

# 勇----------------------------------------------------------------------------------------------------
@ app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == "POST":
        request_data = json.loads(request.data)
        username = request_data['username']

        print(username)
        password = request_data['password']
        login_user = collection.find_one({"username": username})
        print(login_user)
        if login_user:
            print(login_user['password'])
            if bcrypt.checkpw(password.encode('utf-8'), login_user['password']):
                session['username'] = username
                print("login success")
                print(username)
                success = {'message': 'login success', 'whoisLogin': username, '_id': login_user['_id'], 'wallet': login_user['wallet'], 'name1': login_user[
                    'name1'], 'name2': login_user['name2'], 'address': login_user['address'], 'phone': login_user['phone'], 'img': login_user['img'], 'email': login_user['email']}
                print(jsonify(success))
                return jsonify(success)
            else:
                print("Password input error")
                error = {'message': 'Password input error'}
                return jsonify(error)
        else:
            print("Account input error")
            error = {'message': 'Account input error'}
            return jsonify(error)
    return render_template('login.html')


@ app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    else:
        if request.method == 'POST':
            if request.values['send'] == 'logout':
                # session.pop('user_id', None)
                session.clear()
                print('already clear')
                return redirect(url_for('login'))
    return render_template('profile.html')


@ app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == "POST":
        request_data = json.loads(request.data)
        name1 = request_data['name1']
        name2 = request_data['name2']
        username = request_data['username']
        password = request_data['password']
        email = request_data['email']
        postal = request_data['postal']
        address = request_data['address']
        phone = request_data['phone']
        birthday = request_data['birthday']
        sex = request_data['sex']
        existing_user = collection.find_one({"username": username})
        existing_email = collection.find_one({"email": email})
        allcontent = collection.find({})
        maxid = 1
        for x in allcontent:
            if existing_user is None:
                maxid = x['_id']+1  # 如果沒有這個User，把ID +1
            else:
                maxid = x['_id']  # 如果有這個User，停在最後一個的ID
        # wallet = hashlib.sha256(
        #     f'{username}+{maxid}'.encode('utf-8')).hexdigest()  # 加密使用者帳號，產生錢包位址
        wallet = u''.join(generator.choice(alphabet) for _ in range(81))
        correctaddress = False
        for i in data:
            if i["郵遞區號"] == str(postal) and i["行政區名"] == address[0:6]:
                correctaddress = True
                print("True address")
            else:
                erroraddress = {'message': 'The address is wrong'}
                print("error address")
                print(jsonify(erroraddress))
        print(correctaddress)
        if existing_user is None:
            if correctaddress is True:
                if existing_email is None:
                    hashpassword = bcrypt.hashpw(
                        password.encode('utf-8'), bcrypt.gensalt())  # 加密密碼
                    collection.insert_one({'_id': maxid, 'wallet': wallet, 'name1': name1, 'name2': name2, 'username': username,
                                           'password': hashpassword, 'email': email, 'address': address, 'phone': phone, 'birthday': birthday, 'sex': sex, 'img': "https://img.tukuppt.com/png_preview/00/04/40/bbxGjzVlLI.jpg!/fw/780"})
                    session['username'] = username
                    success = {'message': 'insert success'}
                    blockchain.clientRegister(maxid)
                    print(jsonify(success))
                    return jsonify(success)
                error = {'message': 'That email already exists'}
                print(jsonify(error))
                return jsonify(error)
            erroraddress = {'message': 'The address is wrong'}
            print(jsonify(erroraddress))
            return(jsonify(erroraddress))
        error = {'message': 'That username already exists'}
        print(jsonify(error))
        return jsonify(error)
    return render_template('register.html')


@ app.route('/updatecustomerimage', methods=['GET', 'POST'])
def updateimage():
    if request.method == "POST":
        request_data = json.loads(request.data)
        username = request_data['username']
        print(username)
        image = request_data['image']
        print(image)

        collection.update({'username': username}, {"$set": {'img': image}})
        # collection2.update({'_id': 0}, {"$set": {'P_Image': icon}})
        # maxid = 0
        # for x in allcontent:
        #     maxid = x['_id']+1
        # collection3.insert({"_id":maxid, "filename":filename, "img": icon})
        # project = {"img":True}
        # a = collection2.find_one({"_id":2} , projection=project)
        # print(a['img'])
        successupdate = {'message': 'The image is success upload'}
        print("successupdate")
    return jsonify("success")


@ app.route('/project', methods=['GET', 'POST'])
def project():
    if request.method == "POST":
        request_data = json.loads(request.data)
        filename = request_data['filename']
        icon = request_data['icon']
        successupload = {'message': 'The image is success upload'}
        allcontent = collection3.find({})
        maxid = 0
        for x in allcontent:
            maxid = x['_id']+1
        collection3.insert({"_id": maxid, "filename": filename, "img": icon})
        # project = {"img":True}
        # a = collection2.find_one({"_id":2} , projection=project)
        # print(a['img'])
        print("success")
    return jsonify("success")


@ app.route('/updatelikeitem', methods=['GET', 'POST'])
def updatelikeitem():
    if request.method == "POST":
        request_data = json.loads(request.data)
        id = request_data['id']
        like = request_data['like']
        successupload = {'message': 'Like item is success upload'}
        allcontent = collection.find({})
        for x in allcontent:
            collection.update({'_id': id}, {"$set": {'like': like}})
        # project = {"img":True}
        # a = collection2.find_one({"_id":2} , projection=project)
        # print(a['img'])
        print("success")
    return jsonify(successupload)

@ app.route('/InsertProject', methods=['GET', 'POST'])
def InsertProject():
    if request.method == "POST":
        request_data = json.loads(request.data)
        wallet=request_data["wallet"]
        P_Name = request_data['P_Name']
        P_Address = request_data['P_Address']
        postal=request_data['postal']
        P_Content1 = request_data['P_Content1']
        P_Content2 = request_data['P_Content2']
        P_Image=request_data['P_Image']
        P_ContentImage= request_data['P_ContentImage']
        P_Date = request_data['P_Date']
        type = request_data['type']
        capacity = request_data['capacity']
        ping = request_data['ping']
        price = request_data['price']
        username = request_data['username']
        existing_Name = collection2.find_one({"P_Name": P_Name})
 
        allcontent = collection2.find({})
        maxid = 1
        for x in allcontent:
            if existing_Name is None:
                maxid = x['_id']+1  # 如果沒有這個User，把ID +1
            else:
                maxid = x['_id']  # 如果有這個User，停在最後一個的ID
        # wallet = hashlib.sha256(
        #     f'{username}+{maxid}'.encode('utf-8')).hexdigest()  # 加密使用者帳號，產生錢包位址
        print("郵遞區號")
        print(postal)
        print(P_Address)
        correctaddress = False
        for i in data:
            if i["郵遞區號"] == str(postal) and i["行政區名"] == P_Address[0:6]:
                correctaddress = True
                print("True address")
                break
            else:
                erroraddress = {'message': 'The address is wrong'}
               
        print(correctaddress)
 
        if existing_Name is None:
            if correctaddress is True:                    
                collection2.insert_one({'_id': maxid, 'wallet': wallet, 'P_Name': P_Name, 'P_Address': P_Address,
                                           'P_Content1': P_Content1, 'P_Content2': P_Content2, 'P_Image':P_Image, 'P_ContentImage':P_ContentImage, 'P_Date': P_Date, 'type': type, 'capacity': capacity, 'ping': ping,'price':price ,'username':username})
                success = {'message': 'insert success'}                                
                print(jsonify(success))
                return jsonify(success)
            erroraddress = {'message': "The address is wrong"}
            print(jsonify(erroraddress))
            return(jsonify(erroraddress))
        error = {'message': 'That P_Name already exist'}      
        print(jsonify(error))
        return jsonify(error)
    
if __name__ == "__main__":
    app.run(debug=True)
